import os
from pathlib import Path

# PDF
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# DOCX
try:
    from docx import Document
except ImportError:
    Document = None

# Images / OCR
try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None


# ============================================================
# FILE TYPE DETECTION
# ============================================================

def get_file_type(filename):
    """
    Determine whether a file is structured data
    or an unstructured document.
    """

    extension = Path(filename).suffix.lower()

    structured_extensions = {
        ".csv",
        ".xlsx",
        ".xls",
    }

    document_extensions = {
        ".pdf",
        ".docx",
        ".doc",
        ".txt",
        ".md",
        ".jpg",
        ".jpeg",
        ".png",
        ".webp",
    }

    if extension in structured_extensions:
        return "structured"

    if extension in document_extensions:
        return "document"

    return "unknown"


# ============================================================
# PDF EXTRACTION
# ============================================================

def extract_pdf_text(file_path):
    """
    Extract text from a PDF file.
    """

    if PdfReader is None:
        raise ImportError(
            "pypdf is not installed. "
            "Run: pip install pypdf"
        )

    reader = PdfReader(file_path)

    pages = []

    for page_number, page in enumerate(reader.pages, start=1):

        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""

        text = text.strip()

        if text:
            pages.append(
                f"[Page {page_number}]\n{text}"
            )

    return "\n\n".join(pages)


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx_text(file_path):
    """
    Extract paragraphs and tables from a DOCX file.
    """

    if Document is None:
        raise ImportError(
            "python-docx is not installed. "
            "Run: pip install python-docx"
        )

    document = Document(file_path)

    sections = []

    # Paragraphs
    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            sections.append(text)

    # Tables
    for table in document.tables:

        rows = []

        for row in table.rows:

            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            rows.append(
                " | ".join(cells)
            )

        if rows:
            sections.append(
                "\n".join(rows)
            )

    return "\n\n".join(sections)


# ============================================================
# TEXT FILE EXTRACTION
# ============================================================

def extract_text_file(file_path):
    """
    Extract text from TXT or Markdown files.
    """

    encodings = [
        "utf-8",
        "utf-8-sig",
        "latin-1",
    ]

    for encoding in encodings:

        try:

            with open(
                file_path,
                "r",
                encoding=encoding,
            ) as file:

                return file.read()

        except UnicodeDecodeError:

            continue

    raise UnicodeDecodeError(
        "unknown",
        b"",
        0,
        1,
        "Could not decode text file.",
    )


# ============================================================
# IMAGE OCR
# ============================================================

def extract_image_text(file_path):
    """
    Extract text from JPG, JPEG, PNG or WEBP
    using OCR.
    """

    if Image is None:
        raise ImportError(
            "Pillow is not installed. "
            "Run: pip install pillow"
        )

    if pytesseract is None:
        raise ImportError(
            "pytesseract is not installed. "
            "Run: pip install pytesseract"
        )

    image = Image.open(file_path)

    text = pytesseract.image_to_string(
        image
    )

    return text.strip()


# ============================================================
# MAIN DOCUMENT EXTRACTION
# ============================================================

def extract_document_text(file_path):
    """
    Automatically extract text based on file extension.
    """

    extension = Path(file_path).suffix.lower()

    if extension == ".pdf":

        return extract_pdf_text(
            file_path
        )

    if extension == ".docx":

        return extract_docx_text(
            file_path
        )

    if extension in {
        ".txt",
        ".md",
    }:

        return extract_text_file(
            file_path
        )

    if extension in {
        ".jpg",
        ".jpeg",
        ".png",
        ".webp",
    }:

        return extract_image_text(
            file_path
        )

    raise ValueError(
        f"Unsupported document type: {extension}"
    )


# ============================================================
# DOCUMENT INFORMATION
# ============================================================

def get_document_info(file_path):
    """
    Return basic information about a document.
    """

    path = Path(file_path)

    file_type = get_file_type(
        path.name
    )

    return {
        "filename": path.name,
        "extension": path.suffix.lower(),
        "file_type": file_type,
        "size_bytes": path.stat().st_size,
    }


# ============================================================
# DOCUMENT LOADER
# ============================================================

def load_document(file_path):
    """
    Complete document loading pipeline.

    Returns:
        {
            filename,
            file_type,
            extension,
            text,
            characters
        }
    """

    path = Path(file_path)

    if not path.exists():

        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    file_type = get_file_type(
        path.name
    )

    if file_type != "document":

        raise ValueError(
            f"{path.name} is not a document file."
        )

    text = extract_document_text(
        file_path
    )

    return {
        "filename": path.name,
        "file_type": file_type,
        "extension": path.suffix.lower(),
        "text": text,
        "characters": len(text),
    }


# ============================================================
# DOCUMENT CHUNKING
# ============================================================

def chunk_text(
    text,
    chunk_size=1500,
    overlap=200,
):
    """
    Split document text into overlapping chunks.

    This will later be used for document AI/RAG.
    """

    if not text:
        return []

    text = str(text).strip()

    if not text:
        return []

    if overlap >= chunk_size:

        raise ValueError(
            "overlap must be smaller than chunk_size"
        )

    chunks = []

    start = 0
    text_length = len(text)

    while start < text_length:

        end = min(
            start + chunk_size,
            text_length,
        )

        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        if end >= text_length:
            break

        start = end - overlap

    return chunks


# ============================================================
# DOCUMENT PROCESSOR
# ============================================================

def process_document(
    file_path,
    chunk_size=1500,
    overlap=200,
):
    """
    Load a document, extract its text,
    and split it into chunks.
    """

    document = load_document(
        file_path
    )

    chunks = chunk_text(
        document["text"],
        chunk_size=chunk_size,
        overlap=overlap,
    )

    document["chunks"] = chunks

    document["chunk_count"] = len(
        chunks
    )

    return document