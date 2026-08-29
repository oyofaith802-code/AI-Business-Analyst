# ============================================================
# AI BUSINESS ANALYST - UNIVERSAL DATASET UPLOAD
# ============================================================

import os
import re
import json

import pandas as pd

from sqlalchemy import text

from database import engine

from schema_memory import (
    save_existing_table_schema
)

from dataset_profile_storage import (
    save_dataset_profile
)

from plan_manager import (
    upload_allowed,
    record_upload
)


# ============================================================
# SUPPORTED FILE TYPES
# ============================================================

SUPPORTED_EXTENSIONS = {
    ".csv",
    ".xlsx",
    ".xls",
    ".json",
    ".parquet",
    ".pdf",
    ".docx",
}


# ============================================================
# CLEAN TABLE NAME
# ============================================================

def clean_table_name(filename):

    name = os.path.basename(str(filename))

    name = os.path.splitext(name)[0]

    name = name.lower()

    name = re.sub(
        r"[^a-z0-9_]+",
        "_",
        name
    )

    name = re.sub(
        r"_+",
        "_",
        name
    )

    name = name.strip("_")

    if not name:
        name = "uploaded_dataset"

    return name[:55]


# ============================================================
# CLEAN COLUMN NAMES
# ============================================================

def clean_column_names(df):

    columns = []

    used = {}

    for column in df.columns:

        column = str(column).strip().lower()

        column = re.sub(
            r"[^a-z0-9_]+",
            "_",
            column
        )

        column = re.sub(
            r"_+",
            "_",
            column
        )

        column = column.strip("_")

        if not column:

            column = "column"

        if column in used:

            used[column] += 1

            column = (
                f"{column}_{used[column]}"
            )

        else:

            used[column] = 0

        columns.append(column)

    df.columns = columns

    return df


# ============================================================
# DETECT DATES
# ============================================================

def detect_dates(df):

    for column in df.columns:

        if df[column].dtype != "object":

            continue

        try:

            converted = pd.to_datetime(
                df[column],
                errors="coerce"
            )

            valid_ratio = (
                converted.notna().mean()
            )

            if valid_ratio >= 0.8:

                df[column] = converted

        except Exception:

            pass

    return df


# ============================================================
# READ CSV
# ============================================================

def read_csv_file(file):

    try:

        file.seek(0)

        return pd.read_csv(file)

    except UnicodeDecodeError:

        file.seek(0)

        return pd.read_csv(
            file,
            encoding="latin1"
        )


# ============================================================
# READ EXCEL
# ============================================================

def read_excel_file(file):

    file.seek(0)

    return pd.read_excel(file)


# ============================================================
# READ JSON
# ============================================================

def read_json_file(file):

    file.seek(0)

    data = json.load(file)

    if isinstance(data, list):

        return pd.DataFrame(data)

    if isinstance(data, dict):

        try:

            return pd.json_normalize(data)

        except Exception:

            return pd.DataFrame([data])

    raise ValueError(
        "Unsupported JSON structure."
    )


# ============================================================
# READ PARQUET
# ============================================================

def read_parquet_file(file):

    file.seek(0)

    return pd.read_parquet(file)


# ============================================================
# READ PDF
# ============================================================

def read_pdf_file(file):

    try:

        from pypdf import PdfReader

    except ImportError:

        raise ImportError(
            "Install pypdf with: pip install pypdf"
        )

    file.seek(0)

    reader = PdfReader(file)

    rows = []

    for page_number, page in enumerate(
        reader.pages,
        start=1
    ):

        page_text = page.extract_text()

        if not page_text:

            continue

        lines = [
            line.strip()
            for line in page_text.splitlines()
            if line.strip()
        ]

        for line in lines:

            rows.append(
                {
                    "page": page_number,
                    "content": line
                }
            )

    if not rows:

        raise ValueError(
            "No readable text was found in the PDF."
        )

    return pd.DataFrame(rows)


# ============================================================
# READ WORD
# ============================================================

def read_word_file(file):

    try:

        from docx import Document

    except ImportError:

        raise ImportError(
            "Install python-docx with: pip install python-docx"
        )

    file.seek(0)

    document = Document(file)

    rows = []

    # --------------------------------------------------------
    # PARAGRAPHS
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        content = paragraph.text.strip()

        if content:

            rows.append(
                {
                    "source": "paragraph",
                    "content": content
                }
            )

    # --------------------------------------------------------
    # TABLES
    # --------------------------------------------------------

    for table_number, table in enumerate(
        document.tables,
        start=1
    ):

        for row in table.rows:

            values = [
                cell.text.strip()
                for cell in row.cells
            ]

            rows.append(
                {
                    "source":
                        f"table_{table_number}",

                    "content":
                        " | ".join(values)
                }
            )

    if not rows:

        raise ValueError(
            "No readable content was found in the Word document."
        )

    return pd.DataFrame(rows)


# ============================================================
# READ UPLOADED FILE
# ============================================================

def read_uploaded_file(uploaded_file):

    filename = str(
        uploaded_file.name
    )

    extension = os.path.splitext(
        filename
    )[1].lower()

    if extension not in SUPPORTED_EXTENSIONS:

        raise ValueError(
            f"Unsupported file type: {extension}"
        )

    if extension == ".csv":

        return read_csv_file(
            uploaded_file
        )

    if extension in (
        ".xlsx",
        ".xls"
    ):

        return read_excel_file(
            uploaded_file
        )

    if extension == ".json":

        return read_json_file(
            uploaded_file
        )

    if extension == ".parquet":

        return read_parquet_file(
            uploaded_file
        )

    if extension == ".pdf":

        return read_pdf_file(
            uploaded_file
        )

    if extension == ".docx":

        return read_word_file(
            uploaded_file
        )

    raise ValueError(
        "Unsupported file format."
    )


# ============================================================
# CREATE DATASET PROFILE
# ============================================================

def create_dataset_profile(df):

    profile = {

        "rows":
            int(len(df)),

        "columns":
            int(len(df.columns)),

        "column_names":
            list(df.columns),

        "column_profiles":
            [],

        "numeric_summary":
            {}

    }

    for column in df.columns:

        series = df[column]

        info = {

            "column":
                column,

            "dtype":
                str(series.dtype),

            "missing":
                int(series.isna().sum()),

            "unique":
                int(series.nunique())

        }

        if pd.api.types.is_numeric_dtype(
            series
        ):

            clean = series.dropna()

            if not clean.empty:

                info["minimum"] = float(
                    clean.min()
                )

                info["maximum"] = float(
                    clean.max()
                )

                info["average"] = float(
                    clean.mean()
                )

        profile[
            "column_profiles"
        ].append(info)

    # --------------------------------------------------------
    # NUMERIC SUMMARY
    # --------------------------------------------------------

    numeric_columns = df.select_dtypes(
        include="number"
    ).columns

    for column in numeric_columns:

        series = df[column].dropna()

        if series.empty:

            continue

        profile[
            "numeric_summary"
        ][column] = {

            "sum":
                float(series.sum()),

            "average":
                float(series.mean()),

            "minimum":
                float(series.min()),

            "maximum":
                float(series.max())

        }

    return profile


# ============================================================
# CHECK TABLE EXISTS
# ============================================================

def table_exists(table_name):

    sql = text(
        """
        SELECT EXISTS (
            SELECT 1
            FROM information_schema.tables
            WHERE table_schema = 'public'
            AND table_name = :table_name
        )
        """
    )

    with engine.connect() as conn:

        return conn.execute(
            sql,
            {
                "table_name":
                    table_name
            }
        ).scalar()


# ============================================================
# CREATE UNIQUE TABLE NAME
# ============================================================

def create_unique_table_name(base_name):

    table_name = base_name

    counter = 1

    while table_exists(table_name):

        table_name = (
            f"{base_name}_{counter}"
        )

        counter += 1

    return table_name


# ============================================================
# SAVE DATASET
# ============================================================

def save_dataset(
    df,
    table_name,
    user_email
):

    try:

        df.to_sql(
            table_name,
            engine,
            if_exists="fail",
            index=False
        )

        print(
            f"✅ Dataset saved: {table_name}"
        )

        # ----------------------------------------------------
        # SAVE SCHEMA
        # ----------------------------------------------------

        try:

            save_existing_table_schema(
                table_name
            )

            print(
                "✅ Schema saved."
            )

        except Exception as e:

            print(
                f"⚠️ Schema warning: {e}"
            )

        # ----------------------------------------------------
        # SAVE PROFILE
        # ----------------------------------------------------

        try:

            profile = create_dataset_profile(
                df
            )

            save_dataset_profile(
                user_email=user_email,
                dataset_name=table_name,
                profile=profile
            )

            print(
                "✅ Dataset profile saved."
            )

        except Exception as e:

            print(
                f"⚠️ Profile warning: {e}"
            )

        return True

    except Exception as e:

        print(
            f"❌ Database save failed: {e}"
        )

        return False


# ============================================================
# MAIN UPLOAD PIPELINE
# ============================================================

def process_upload(
    uploaded_file,
    user_email
):

    # --------------------------------------------------------
    # VALIDATE USER
    # --------------------------------------------------------

    if not user_email:

        return {

            "success": False,

            "error":
                "User email is required."

        }

    user_email = (
        str(user_email)
        .strip()
        .lower()
    )

    # --------------------------------------------------------
    # CHECK UPLOAD PLAN LIMIT
    # --------------------------------------------------------

    try:

        allowed = upload_allowed(
            user_email
        )

    except Exception as e:

        print(
            f"❌ Could not check upload limit: {e}"
        )

        return {

            "success": False,

            "error":
                "Could not verify your upload limit."

        }

    if not allowed:

        return {

            "success": False,

            "error":
                "You have reached your monthly upload limit for your current plan."

        }

    # --------------------------------------------------------
    # VALIDATE FILE
    # --------------------------------------------------------

    if uploaded_file is None:

        return {

            "success": False,

            "error":
                "No file was selected."

        }

    filename = str(
        uploaded_file.name
    )

    extension = os.path.splitext(
        filename
    )[1].lower()

    if extension not in SUPPORTED_EXTENSIONS:

        return {

            "success": False,

            "error":
                f"Unsupported file type: {extension}"

        }

    # --------------------------------------------------------
    # READ FILE
    # --------------------------------------------------------

    try:

        df = read_uploaded_file(
            uploaded_file
        )

    except Exception as e:

        return {

            "success": False,

            "error":
                f"Could not read the file: {e}"

        }

    # --------------------------------------------------------
    # VALIDATE DATA
    # --------------------------------------------------------

    if df is None or df.empty:

        return {

            "success": False,

            "error":
                "The file contains no usable data."

        }

    if len(df.columns) == 0:

        return {

            "success": False,

            "error":
                "No usable columns were found."

        }

    # --------------------------------------------------------
    # CLEAN DATA
    # --------------------------------------------------------

    try:

        df = clean_column_names(df)

        df = detect_dates(df)

    except Exception as e:

        return {

            "success": False,

            "error":
                f"Could not clean the dataset: {e}"

        }

    # --------------------------------------------------------
    # CREATE TABLE NAME
    # --------------------------------------------------------

    base_name = clean_table_name(
        filename
    )

    table_name = create_unique_table_name(
        base_name
    )

    # --------------------------------------------------------
    # SAVE DATASET
    # --------------------------------------------------------

    success = save_dataset(
        df,
        table_name,
        user_email
    )

    if not success:

        return {

            "success": False,

            "error":
                "Failed to save dataset."

        }

    # --------------------------------------------------------
    # RECORD UPLOAD USAGE
    # --------------------------------------------------------

    try:

        record_upload(
            user_email
        )

        print(
            "✅ Upload usage recorded."
        )

    except Exception as e:

        print(
            f"⚠️ Upload saved, but usage could not be recorded: {e}"
        )

    # --------------------------------------------------------
    # RETURN RESULT
    # --------------------------------------------------------

    return {

        "success":
            True,

        "table_name":
            table_name,

        "rows":
            int(len(df)),

        "columns":
            int(len(df.columns)),

        "column_names":
            list(df.columns),

        "file_type":
            extension

    }


# ============================================================
# TEST
# ============================================================

if __name__ == "__main__":

    print(
        "============================================================"
    )

    print(
        "AI BUSINESS ANALYST - DATASET UPLOAD"
    )

    print(
        "============================================================"
    )

    print()

    print(
        "Universal dataset upload module loaded successfully."
    )

    print()

    print(
        "Supported formats:"
    )

    for extension in sorted(
        SUPPORTED_EXTENSIONS
    ):

        print(
            f"• {extension}"
        )

    print()

    print(
        "✅ Upload module ready."
    )
